from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from flask_jwt_extended import JWTManager, create_access_token, jwt_required, get_jwt_identity, get_jwt
import os
from werkzeug.utils import secure_filename
import PyPDF2
from docx import Document
import requests
from dotenv import load_dotenv
import json
from datetime import timedelta
import time

load_dotenv()

app = Flask(__name__)
CORS(app, resources={
    r"/api/*": {
        "origins": ["http://localhost:5174"],
        "methods": ["GET", "POST", "PUT", "DELETE", "OPTIONS"],
        "allow_headers": ["Content-Type", "Authorization"],
        "supports_credentials": True
    }
})

# 配置
app.config['JWT_SECRET_KEY'] = os.getenv('JWT_SECRET_KEY', 'your-secret-key')
app.config['JWT_ACCESS_TOKEN_EXPIRES'] = timedelta(minutes=30)  # 设置 token 过期时间为 30 分钟
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

# DeepSeek配置
DEEPSEEK_API_KEY = os.getenv('DEEPSEEK_API_KEY')
DEEPSEEK_API_URL = "https://api.deepseek.com/v1/chat/completions"

# 文件解析状态存储
parsing_status = {}

# 标书生成状态存储
generation_status = {}

# 存储生成结果的字典
generation_results = {}

jwt = JWTManager(app)

# 模拟用户数据
USERS = {
    'admin': 'admin123'
}

def analyze_tender_content(content: str) -> dict:
    """使用 DeepSeek 分析招标文件内容"""
    prompt = f"""
    请分析以下招标文件内容，提取关键信息、评分标准并生成标书大纲。请以JSON格式返回，格式如下：
    {{
        "project_info": {{
            "name": "项目名称",
            "type": "项目类型",
            "budget": "预算金额",
            "deadline": "截止日期",
            "requirements": ["主要要求1", "主要要求2", ...]
        }},
        "scoring_criteria": [
            {{
                "id": "1",
                "category": "评分类别",
                "item": "评分项",
                "score": "分值",
                "description": "评分标准描述",
                "requirements": ["具体要求1", "具体要求2", ...]
            }},
            ...
        ],
        "outline": [
            {{
                "id": "1",
                "title": "章节标题",
                "required": true/false,
                "description": "章节描述",
                "key_points": ["关键点1", "关键点2", ...]
            }},
            ...
        ]
    }}

    招标文件内容：
    {content}
    """

    try:
        print("开始调用 DeepSeek API...")
        print(f"API Key: {DEEPSEEK_API_KEY[:5]}...")  # 只打印前5位
        print(f"API URL: {DEEPSEEK_API_URL}")
        
        headers = {
            "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
            "Content-Type": "application/json"
        }
        
        data = {
            "model": "deepseek-chat",
            "messages": [
                {"role": "system", "content": "你是一个专业的标书分析专家，擅长从招标文件中提取关键信息、评分标准并生成标书大纲。请始终以有效的JSON格式返回结果。"},
                {"role": "user", "content": prompt}
            ],
            "temperature": 0.3,
            "response_format": {"type": "json_object"}
        }
        
        print("发送请求数据:", json.dumps(data, ensure_ascii=False)[:200] + "...")
        
        response = requests.post(DEEPSEEK_API_URL, headers=headers, json=data)
        print(f"API 响应状态码: {response.status_code}")
        
        response.raise_for_status()
        
        response_json = response.json()
        print(f"API 响应内容: {response_json['choices'][0]['message']['content'][:200]}...")
        
        # 检查返回内容是否为空
        content_text = response_json['choices'][0]['message']['content']
        if not content_text or content_text.isspace():
            raise ValueError("API 返回内容为空")
            
        try:
            result = json.loads(content_text)
            # 验证结果格式
            if not isinstance(result, dict):
                raise ValueError("返回结果不是有效的 JSON 对象")
                
            # 确保至少包含基本字段
            if 'project_info' not in result or 'outline' not in result:
                raise ValueError("返回结果缺少必要字段")
                
            return result
        except json.JSONDecodeError as json_err:
            print(f"JSON 解析错误: {json_err}")
            print(f"尝试解析的内容: {content_text[:500]}...")
            raise
    except Exception as e:
        print(f"分析错误: {str(e)}")
        # 提供默认返回值
        default_result = {
            "project_info": {
                "name": "未识别到项目名称",
                "type": "未知",
                "budget": "未知",
                "deadline": "未知",
                "requirements": ["需求提取失败"]
            },
            "scoring_criteria": [
                {
                    "id": "1",
                    "category": "未能识别评分类别",
                    "item": "默认评分项",
                    "score": "100",
                    "description": "无法从文档中提取评分标准",
                    "requirements": ["请手动编辑评分要求"]
                }
            ],
            "outline": generate_standard_outline()
        }
        
        # 确保每个大纲项都有 key_points 字段
        for item in default_result["outline"]:
            if "key_points" not in item:
                item["key_points"] = ["请填写关键点"]
                
        return default_result

@app.route('/api/login', methods=['POST'])
def login():
    data = request.get_json()
    username = data.get('username')
    password = data.get('password')
    
    if username in USERS and USERS[username] == password:
        access_token = create_access_token(identity=username)
        return jsonify({'token': access_token}), 200
    
    return jsonify({'error': 'Invalid credentials'}), 401

@app.route('/api/upload', methods=['POST'])
@jwt_required()
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)
        
        # 创建任务ID并初始化解析状态
        task_id = f"parse_{int(time.time())}"
        parsing_status[task_id] = {
            'progress': 10,
            'status': 'processing',
            'message': '文件已上传，开始解析...'
        }
        
        # 开始解析过程
        try:
            # 更新状态
            parsing_status[task_id]['progress'] = 30
            parsing_status[task_id]['message'] = '正在解析文件内容...'
            
            # 解析文件内容
            content = parse_document(filepath)
            
            # 更新状态
            parsing_status[task_id]['progress'] = 60
            parsing_status[task_id]['message'] = '正在使用AI分析内容...'
            
            # 使用大模型分析内容
            analysis_result = analyze_tender_content(content)
            
            # 更新状态
            parsing_status[task_id]['progress'] = 100
            parsing_status[task_id]['status'] = 'success'
            parsing_status[task_id]['message'] = '解析完成'
            
            return jsonify({
                'message': '文件上传成功',
                'task_id': task_id,
                'content': content,
                'project_info': analysis_result['project_info'],
                'scoring_criteria': analysis_result['scoring_criteria'],
                'outline': analysis_result['outline']
            }), 200
        except Exception as e:
            # 更新状态
            parsing_status[task_id]['status'] = 'error'
            parsing_status[task_id]['message'] = f'解析失败: {str(e)}'
            return jsonify({'error': f'解析失败: {str(e)}', 'task_id': task_id}), 500
    
    return jsonify({'error': 'Invalid file type'}), 400

@app.route('/api/parsing-status/<task_id>', methods=['GET'])
@jwt_required()
def get_parsing_status(task_id):
    """获取文件解析进度"""
    if task_id in parsing_status:
        return jsonify(parsing_status[task_id]), 200
    return jsonify({'error': 'Task not found'}), 404

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in {'pdf', 'docx', 'doc'}

def parse_document(filepath):
    if filepath.endswith('.pdf'):
        return parse_pdf(filepath)
    elif filepath.endswith(('.docx', '.doc')):
        return parse_docx(filepath)
    return None

def parse_pdf(filepath):
    content = []
    with open(filepath, 'rb') as file:
        pdf_reader = PyPDF2.PdfReader(file)
        for page in pdf_reader.pages:
            content.append(page.extract_text())
    return '\n'.join(content)

def parse_docx(filepath):
    doc = Document(filepath)
    return '\n'.join([paragraph.text for paragraph in doc.paragraphs])

def extract_project_name(content: str) -> str:
    # 简单的项目名称提取逻辑
    # 在实际应用中，这里应该使用更复杂的NLP处理
    lines = content.split('\n')
    for line in lines:
        if '项目名称' in line or '工程名称' in line:
            return line.split('：')[-1].strip()
    return '未识别到项目名称'

def generate_standard_outline() -> list:
    return [
        {
            'id': '1',
            'title': '投标函',
            'required': True,
            'description': '投标人对招标文件的响应和承诺'
        },
        {
            'id': '2',
            'title': '投标报价表',
            'required': True,
            'description': '详细的报价信息'
        },
        {
            'id': '3',
            'title': '技术方案',
            'required': True,
            'description': '详细的技术实施方案'
        },
        {
            'id': '4',
            'title': '项目团队',
            'required': True,
            'description': '项目团队成员介绍'
        },
        {
            'id': '5',
            'title': '业绩证明',
            'required': True,
            'description': '类似项目业绩证明'
        },
        {
            'id': '6',
            'title': '资质证书',
            'required': True,
            'description': '相关资质证书'
        },
        {
            'id': '7',
            'title': '售后服务方案',
            'required': False,
            'description': '售后服务承诺和方案'
        },
        {
            'id': '8',
            'title': '其他补充材料',
            'required': False,
            'description': '其他补充说明材料'
        }
    ]

@app.route('/api/refresh', methods=['POST'])
@jwt_required(refresh=True)
def refresh():
    current_user = get_jwt_identity()
    new_token = create_access_token(identity=current_user)
    return jsonify({'token': new_token}), 200

@app.route('/api/generate-proposal', methods=['POST'])
@jwt_required()
def generate_proposal():
    """生成标书内容"""
    data = request.get_json()
    outline = data.get('outline', [])
    project_info = data.get('projectInfo', {})
    scoring_criteria = data.get('scoringCriteria', [])
    task_id = data.get('task_id')
    
    if not task_id:
        return jsonify({'error': '缺少任务ID'}), 400
    
    # 初始化生成状态
    update_generation_status(task_id, 0, 'processing', '开始生成标书...')
    
    try:
        # 构建提示词
        prompt = f"""
        请根据以下信息生成一份完整的标书内容。要求语言专业、结构清晰、内容详实。
        
        项目信息：
        项目名称：{project_info.get('name', '未知项目')}
        项目类型：{project_info.get('type', '未知')}
        预算金额：{project_info.get('budget', '未知')}
        截止日期：{project_info.get('deadline', '未知')}
        
        项目要求：
        {', '.join(project_info.get('requirements', ['无具体要求']))}
        
        评分标准：
        """
        
        for criterion in scoring_criteria:
            prompt += f"\n{criterion.get('category', '')} - {criterion.get('item', '')}: {criterion.get('description', '')}"
            if criterion.get('requirements'):
                prompt += f"\n具体要求: {', '.join(criterion.get('requirements', []))}"
        
        prompt += "\n\n请按照以下大纲生成标书内容（以HTML格式输出，便于前端显示）：\n"
        
        for item in outline:
            prompt += f"\n# {item.get('title')}\n"
            prompt += f"{item.get('description')}\n"
            prompt += f"关键点: {', '.join(item.get('key_points', []))}\n"
        
        # 更新状态
        update_generation_status(task_id, 30, 'processing', '正在生成标书内容...')
        
        # 调用AI生成标书内容
        headers = {
            "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
            "Content-Type": "application/json"
        }
        
        data = {
            "model": "deepseek-chat",
            "messages": [
                {
                    "role": "system", 
                    "content": "你是一位专业的标书撰写专家，擅长根据项目要求和评分标准生成专业、详实的标书内容。"
                },
                {"role": "user", "content": prompt}
            ],
            "temperature": 0.5
        }
        
        # 更新状态
        update_generation_status(task_id, 60, 'processing', '正在调用AI模型...')
        
        response = requests.post(DEEPSEEK_API_URL, headers=headers, json=data)
        response.raise_for_status()
        
        # 更新状态
        update_generation_status(task_id, 90, 'processing', '正在处理生成结果...')
        
        # 处理响应结果
        response_json = response.json()
        proposal_content = response_json['choices'][0]['message']['content']
        
        # 将生成的内容转为HTML格式
        html_content = format_proposal_content(proposal_content)
        
        # 存储生成结果
        generation_results[task_id] = html_content
        
        # 更新状态
        update_generation_status(task_id, 100, 'success', '标书生成完成')
        
        return jsonify({'content': html_content}), 200
    except Exception as e:
        print(f"生成标书错误: {str(e)}")
        update_generation_status(task_id, 0, 'error', f'生成失败: {str(e)}')
        return jsonify({'error': '生成标书失败'}), 500

def format_proposal_content(content):
    """将Markdown格式转换为HTML格式"""
    # 简单的格式转换逻辑，实际项目中可以使用Markdown库
    lines = content.split('\n')
    html_lines = []
    
    for line in lines:
        if line.startswith('# '):
            html_lines.append(f'<h1>{line[2:]}</h1>')
        elif line.startswith('## '):
            html_lines.append(f'<h2>{line[3:]}</h2>')
        elif line.startswith('### '):
            html_lines.append(f'<h3>{line[4:]}</h3>')
        elif line.startswith('- '):
            html_lines.append(f'<li>{line[2:]}</li>')
        else:
            html_lines.append(f'<p>{line}</p>')
    
    return ''.join(html_lines)

@app.route('/api/save-proposal', methods=['POST'])
@jwt_required()
def save_proposal():
    """保存标书内容"""
    data = request.get_json()
    content = data.get('content', '')
    project_info = data.get('projectInfo', {})
    project_name = project_info.get('name', 'unnamed_project')
    
    # 创建保存目录
    save_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'proposals')
    os.makedirs(save_dir, exist_ok=True)
    
    # 保存HTML文件
    file_path = os.path.join(save_dir, f"{project_name}_{int(time.time())}.html")
    with open(file_path, 'w', encoding='utf-8') as f:
        f.write(content)
    
    return jsonify({'message': '标书保存成功', 'path': file_path}), 200

@app.route('/api/download-proposal', methods=['POST'])
@jwt_required()
def download_proposal():
    """下载标书为DOCX"""
    data = request.get_json()
    content = data.get('content', '')
    project_info = data.get('projectInfo', {})
    project_name = project_info.get('name', 'unnamed_project')
    
    try:
        # 创建一个临时目录
        temp_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'temp')
        os.makedirs(temp_dir, exist_ok=True)
        
        # 将HTML内容保存为临时文件
        html_path = os.path.join(temp_dir, f"{project_name}_temp.html")
        with open(html_path, 'w', encoding='utf-8') as f:
            f.write(f"""<!DOCTYPE html>
            <html>
            <head>
                <meta charset="utf-8">
                <title>{project_name}</title>
                <style>
                    body {{ font-family: 'SimSun', serif; margin: 2cm; }}
                    h1, h2, h3 {{ color: #333; }}
                    table {{ border-collapse: collapse; width: 100%; }}
                    table, th, td {{ border: 1px solid #ddd; padding: 8px; }}
                </style>
            </head>
            <body>
                {content}
            </body>
            </html>
            """)
        
        # 调用第三方服务或库将HTML转换为DOCX
        # 这里为了简单演示，我们直接使用python-docx
        docx_path = os.path.join(temp_dir, f"{project_name}.docx")
        doc = Document()
        
        # 添加标题
        doc.add_heading(project_name, 0)
        
        # 这里需要解析HTML并添加到docx中
        # 为了简化，我们假设HTML已经被解析成文本
        doc.add_paragraph("项目信息：")
        doc.add_paragraph(f"项目名称：{project_info.get('name', '未知项目')}")
        doc.add_paragraph(f"项目类型：{project_info.get('type', '未知')}")
        doc.add_paragraph(f"预算金额：{project_info.get('budget', '未知')}")
        doc.add_paragraph(f"截止日期：{project_info.get('deadline', '未知')}")
        
        # 导出文档
        doc.save(docx_path)
        
        # 返回文件
        return send_file(
            docx_path,
            as_attachment=True,
            download_name=f"{project_name}.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        print(f"下载标书错误: {str(e)}")
        return jsonify({'error': '下载标书失败'}), 500

@app.route('/api/ai-continue', methods=['POST'])
@jwt_required()
def ai_continue():
    """AI续写功能"""
    data = request.get_json()
    content = data.get('content', '')
    context = data.get('context', {})
    
    prompt = f"""
    请基于下面的标书内容续写，保持风格一致，内容专业：
    
    {content}
    
    请续写关于"{context.get('label', '下一部分')}"的内容。
    """
    
    try:
        headers = {
            "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
            "Content-Type": "application/json"
        }
        
        data = {
            "model": "deepseek-chat",
            "messages": [
                {"role": "system", "content": "你是一位专业的标书撰写专家，擅长续写标书内容。"},
                {"role": "user", "content": prompt}
            ],
            "temperature": 0.3
        }
        
        response = requests.post(DEEPSEEK_API_URL, headers=headers, json=data)
        response.raise_for_status()
        
        response_json = response.json()
        continued_content = response_json['choices'][0]['message']['content']
        
        return jsonify({'continuedContent': continued_content}), 200
    except Exception as e:
        print(f"AI续写错误: {str(e)}")
        return jsonify({'error': 'AI续写失败'}), 500

@app.route('/api/ai-expand', methods=['POST'])
@jwt_required()
def ai_expand():
    """AI扩写功能"""
    data = request.get_json()
    content = data.get('content', '')
    context = data.get('context', {})
    
    prompt = f"""
    请扩展以下标书内容，使其更加详细、专业，增加相关细节和专业术语：
    
    {content}
    
    特别关注"{context.get('label', '全文')}"部分。
    """
    
    try:
        headers = {
            "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
            "Content-Type": "application/json"
        }
        
        data = {
            "model": "deepseek-chat",
            "messages": [
                {"role": "system", "content": "你是一位专业的标书撰写专家，擅长扩展标书内容使其更加专业详实。"},
                {"role": "user", "content": prompt}
            ],
            "temperature": 0.3
        }
        
        response = requests.post(DEEPSEEK_API_URL, headers=headers, json=data)
        response.raise_for_status()
        
        response_json = response.json()
        expanded_content = response_json['choices'][0]['message']['content']
        
        return jsonify({'expandedContent': expanded_content}), 200
    except Exception as e:
        print(f"AI扩写错误: {str(e)}")
        return jsonify({'error': 'AI扩写失败'}), 500

@app.route('/api/ai-polish', methods=['POST'])
@jwt_required()
def ai_polish():
    """AI润色功能"""
    data = request.get_json()
    content = data.get('content', '')
    
    prompt = f"""
    请对以下标书内容进行润色，提升语言表达，修正语法错误，使文档更加专业、流畅：
    
    {content}
    """
    
    try:
        headers = {
            "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
            "Content-Type": "application/json"
        }
        
        data = {
            "model": "deepseek-chat",
            "messages": [
                {"role": "system", "content": "你是一位专业的标书语言专家，擅长润色和优化标书语言表达。"},
                {"role": "user", "content": prompt}
            ],
            "temperature": 0.3
        }
        
        response = requests.post(DEEPSEEK_API_URL, headers=headers, json=data)
        response.raise_for_status()
        
        response_json = response.json()
        polished_content = response_json['choices'][0]['message']['content']
        
        return jsonify({'polishedContent': polished_content}), 200
    except Exception as e:
        print(f"AI润色错误: {str(e)}")
        return jsonify({'error': 'AI润色失败'}), 500

@app.route('/api/image-library', methods=['GET'])
@jwt_required()
def get_image_library():
    """获取图片库"""
    # 创建图片库目录
    image_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'images')
    os.makedirs(image_dir, exist_ok=True)
    
    # 模拟图片数据
    images = [
        {"id": 1, "name": "公司logo", "url": "/static/images/logo.png", "category": "公司形象"},
        {"id": 2, "name": "项目示意图", "url": "/static/images/project.png", "category": "项目展示"},
        {"id": 3, "name": "组织架构", "url": "/static/images/org.png", "category": "团队介绍"},
        {"id": 4, "name": "资质证书", "url": "/static/images/cert.png", "category": "资质证明"},
        {"id": 5, "name": "技术流程", "url": "/static/images/flow.png", "category": "技术方案"}
    ]
    
    return jsonify(images), 200

@app.route('/api/search-content', methods=['POST'])
@jwt_required()
def search_content():
    """联网搜索相关内容"""
    data = request.get_json()
    query = data.get('query', '')
    
    # 模拟搜索结果
    search_results = [
        {
            "title": f"关于{query}的行业标准",
            "content": f"{query}是一种重要的技术/产品/服务，在行业中有明确的标准和规范...",
            "source": "行业标准网"
        },
        {
            "title": f"{query}最新研究进展",
            "content": f"近期研究表明，{query}在应用领域有了新的突破...",
            "source": "技术研究中心"
        },
        {
            "title": f"{query}成功案例分析",
            "content": f"某公司成功应用{query}解决了实际问题...",
            "source": "案例分析网"
        }
    ]
    
    return jsonify(search_results), 200

def update_generation_status(task_id: str, progress: int, status: str, message: str):
    """更新生成状态"""
    generation_status[task_id] = {
        'progress': progress,
        'status': status,
        'message': message
    }

@app.route('/api/generation-status/<task_id>', methods=['GET'])
@jwt_required()
def get_generation_status(task_id):
    """获取标书生成进度"""
    if task_id in generation_status:
        return jsonify(generation_status[task_id]), 200
    return jsonify({'error': 'Task not found'}), 404

@app.route('/api/generation-result/<task_id>', methods=['GET'])
@jwt_required()
def get_generation_result(task_id):
    """获取标书生成结果"""
    if task_id in generation_results:
        return jsonify({'content': generation_results[task_id]}), 200
    return jsonify({'error': 'Task result not found'}), 404

if __name__ == '__main__':
    os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
    static_dir = os.path.join(app.root_path, 'static')
    os.makedirs(static_dir, exist_ok=True)
    app.run(debug=True) 